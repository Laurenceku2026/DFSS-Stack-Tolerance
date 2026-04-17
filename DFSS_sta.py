# app.py
import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from scipy import stats
import math
import re
import base64
from io import BytesIO
from typing import List, Dict, Any, Tuple, Optional
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

st.set_page_config(page_title="Para_Variation - 蒙特卡洛模拟", layout="wide")

st.markdown("""
<style>
    .main-title { font-size: 2.5rem; font-weight: 600; color: #1f3a93; margin-bottom: 1rem; }
    .section-header { font-size: 1.5rem; font-weight: 500; color: #2c3e50; border-left: 5px solid #3498db; padding-left: 15px; margin: 20px 0 15px 0; }
    .metric-card { background-color: #f8f9fa; border-radius: 10px; padding: 15px; text-align: center; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }
    .metric-label { font-size: 1rem; color: #6c757d; margin-bottom: 5px; }
    .metric-value { font-size: 1.8rem; font-weight: 600; color: #2c3e50; }
    .ppm-table { border-collapse: collapse; width: 100%; margin: 0 auto; }
    .ppm-table th, .ppm-table td { border: 2px solid #000000; padding: 10px 16px; text-align: center; font-size: 1rem; }
    .ppm-table th { background-color: #e9ecef; font-weight: 600; }
    .stButton button { background-color: #3498db; color: white; font-weight: 500; border-radius: 5px; font-size: 1.2rem; margin-top: 20px; white-space: pre-line; }
    .stButton button:hover { background-color: #2980b9; }
    .design-value-card { background-color: #e8f4fd; border-radius: 10px; padding: 15px; margin-top: 15px; text-align: center; border-left: 5px solid #3498db; }
    .design-value-card strong { font-size: 1.1rem; }
    .design-value-number { font-size: 1.6rem; font-weight: 600; color: #1f3a93; margin-top: 5px; }
    .big-label { font-size: 1.3rem; font-weight: 500; margin-bottom: 5px; }
    .param-letter { font-weight: bold; font-size: 1rem; text-align: center; background-color: #e9ecef; border-radius: 4px; padding: 6px 0; width: 40px; }
    .formula-hint { font-size: 0.9rem; color: #6c757d; margin-bottom: 5px; }
    .expand-section { background-color: #f8f9fa; border-radius: 8px; padding: 10px; margin-top: 5px; margin-bottom: 10px; border-left: 3px solid #3498db; }
    .sidebar-section { margin-top: 30px; padding-top: 20px; border-top: 1px solid #ddd; }
</style>
""", unsafe_allow_html=True)

# 初始化 session state
if "params" not in st.session_state:
    st.session_state.params = pd.DataFrame({
        "参数名称": ["Cell Cap", "Suction P", "Brush P", "Other(Pump+display)", "V"],
        "均值(Typ)": [2450.0, 70.0, 30.0, 15.0, 3.6],
        "标准差(Std)": [20.74, 0.77, 0.90, 0.45, 0.0036],
        "分布": ["正态分布（完整）", "正态分布（完整）", "正态分布（完整）", "正态分布（完整）", "正态分布（完整）"],
        "分布参数": [{} for _ in range(5)]
    })
if "sim_results_raw" not in st.session_state:
    st.session_state.sim_results_raw = None
if "formula" not in st.session_state:
    st.session_state.formula = "A * E * 7 / 1000 * 60 / (B + C + D)"
if "output_name" not in st.session_state:
    st.session_state.output_name = "Runtime"
if "usl_str" not in st.session_state:
    st.session_state.usl_str = "40.0"
if "lsl_str" not in st.session_state:
    st.session_state.lsl_str = "30.0"
if "analyst_name" not in st.session_state:
    st.session_state.analyst_name = ""
if "analyst_title" not in st.session_state:
    st.session_state.analyst_title = ""

# 分布类型列表
DISTRIBUTIONS = [
    "正态分布（完整）",
    "正态分布（正值）",
    "正态分布（负值）",
    "均匀分布",
    "对数正态分布",
    "威布尔分布",
    "三角分布"
]

def update_param_letters():
    letters = [chr(ord('A') + i) for i in range(len(st.session_state.params))]
    st.session_state.param_letters = {
        row["参数名称"]: letters[i] for i, row in st.session_state.params.iterrows()
    }
update_param_letters()

def parse_limit(s: str) -> Optional[float]:
    if s is None or s.strip() == "":
        return None
    try:
        return float(s)
    except ValueError:
        return None

def sync_usl_from_main(): st.session_state.usl_str = st.session_state.main_usl
def sync_lsl_from_main(): st.session_state.lsl_str = st.session_state.main_lsl
def sync_usl_from_sidebar(): st.session_state.usl_str = st.session_state.usl_sidebar
def sync_lsl_from_sidebar(): st.session_state.lsl_str = st.session_state.lsl_sidebar

def clean_formula(formula: str) -> str:
    formula = formula.strip()
    formula = re.sub(r'\s+', ' ', formula)
    formula = re.sub(r'(?<=[0-9a-zA-Z)])\s*([+\-*/])\s*(?=[0-9a-zA-Z(])', r' \1 ', formula)
    formula = re.sub(r'\(\s+', '(', formula)
    formula = re.sub(r'\s+\)', ')', formula)
    return formula

def replace_letters_with_names(expr: str, param_letters: Dict[str, str]) -> str:
    reverse_map = {v: k for k, v in param_letters.items()}
    for letter, name in sorted(reverse_map.items(), key=lambda x: len(x[0]), reverse=True):
        pattern = r'(?<![a-zA-Z0-9_])' + re.escape(letter) + r'(?![a-zA-Z0-9_])'
        expr = re.sub(pattern, name, expr)
    return expr

def safe_eval_with_mapping(expr: str, param_names: List[str], context_values: List[float], param_letters: Dict[str, str]) -> float:
    expr = clean_formula(expr)
    expr_with_names = replace_letters_with_names(expr, param_letters)
    temp_names = [f"__p{i}__" for i in range(len(param_names))]
    sorted_params = sorted(zip(param_names, temp_names), key=lambda x: len(x[0]), reverse=True)
    expr_temp = expr_with_names
    for orig, temp in sorted_params:
        pattern = r'(?<![a-zA-Z0-9_])' + re.escape(orig) + r'(?![a-zA-Z0-9_])'
        expr_temp = re.sub(pattern, temp, expr_temp)
    context = {temp: val for temp, val in zip(temp_names, context_values)}
    allowed_names = {
        "sqrt": math.sqrt, "exp": math.exp, "log": math.log, "log10": math.log10,
        "sin": math.sin, "cos": math.cos, "tan": math.tan, "pi": math.pi, "e": math.e,
        "abs": abs, "pow": pow
    }
    allowed_names.update(context)
    try:
        result = eval(expr_temp, {"__builtins__": {}}, allowed_names)
        return float(result)
    except Exception:
        return np.nan

def compute_design_value(params_df: pd.DataFrame, formula: str, param_letters: Dict[str, str]) -> Optional[float]:
    param_names = params_df["参数名称"].astype(str).tolist()
    means = params_df["均值(Typ)"].values.astype(float)
    val = safe_eval_with_mapping(formula, param_names, means, param_letters)
    return val if not np.isnan(val) else None

def generate_sample(dist: str, mean: float, std: float, dist_params: Dict, size: int = 1) -> np.ndarray:
    if dist == "正态分布（完整）":
        return np.random.normal(mean, std, size)
    elif dist == "正态分布（正值）":
        a, b = (0 - mean) / std if std > 0 else -np.inf, np.inf
        if std == 0:
            return np.full(size, max(mean, 0))
        return stats.truncnorm.rvs(a, b, loc=mean, scale=std, size=size)
    elif dist == "正态分布（负值）":
        a, b = -np.inf, (0 - mean) / std if std > 0 else np.inf
        if std == 0:
            return np.full(size, min(mean, 0))
        return stats.truncnorm.rvs(a, b, loc=mean, scale=std, size=size)
    elif dist == "均匀分布":
        low = dist_params.get("low", mean - 3*std)
        high = dist_params.get("high", mean + 3*std)
        return np.random.uniform(low, high, size)
    elif dist == "对数正态分布":
        mean_log = dist_params.get("mean_log", 0.0)
        sigma_log = dist_params.get("sigma_log", 1.0)
        return np.random.lognormal(mean_log, sigma_log, size)
    elif dist == "威布尔分布":
        shape = dist_params.get("shape", 1.0)
        scale = dist_params.get("scale", 1.0)
        return np.random.weibull(shape, size) * scale
    elif dist == "三角分布":
        left = dist_params.get("left", mean - 3*std)
        mode = dist_params.get("mode", mean)
        right = dist_params.get("right", mean + 3*std)
        return np.random.triangular(left, mode, right, size)
    else:
        return np.random.normal(mean, std, size)

def run_monte_carlo(params_df: pd.DataFrame, formula: str, n_sim: int, param_letters: Dict[str, str], seed: int = 42) -> Dict[str, Any]:
    np.random.seed(seed)
    n_params = len(params_df)
    param_names = params_df["参数名称"].astype(str).tolist()
    means = params_df["均值(Typ)"].values.astype(float)
    stds = params_df["标准差(Std)"].values.astype(float)
    dists = params_df["分布"].tolist()
    dist_params_list = params_df["分布参数"].tolist()

    samples = np.zeros((n_sim, n_params))
    for i in range(n_params):
        samples[:, i] = generate_sample(dists[i], means[i], stds[i], dist_params_list[i], n_sim)

    results = []
    for i in range(n_sim):
        val = safe_eval_with_mapping(formula, param_names, samples[i, :], param_letters)
        if not np.isnan(val):
            results.append(val)

    results = np.array(results)
    if len(results) == 0:
        st.error("所有公式计算均失败，请检查公式！")
        return None

    mean_out = np.mean(results)
    std_out = np.std(results, ddof=1)
    max_out = np.max(results)
    min_out = np.min(results)

    hist_counts, bin_edges = np.histogram(results, bins=25, density=False)
    bin_centers = (bin_edges[:-1] + bin_edges[1:]) / 2
    x_pdf = np.linspace(min_out, max_out, 200)
    pdf_theory = stats.norm.pdf(x_pdf, mean_out, std_out)

    return {
        "results": results,
        "samples": samples,
        "mean": mean_out,
        "std": std_out,
        "max": max_out,
        "min": min_out,
        "hist_counts": hist_counts,
        "bin_edges": bin_edges,
        "bin_centers": bin_centers,
        "x_pdf": x_pdf,
        "pdf_theory": pdf_theory,
        "param_names": param_names,
    }

def sensitivity_analysis(params_df: pd.DataFrame, formula: str, n_sim: int, param_letters: Dict[str, str], seed: int = 42) -> Tuple[pd.DataFrame, List[float], List[str]]:
    np.random.seed(seed)
    n_params = len(params_df)
    param_names = params_df["参数名称"].astype(str).tolist()
    means = params_df["均值(Typ)"].values.astype(float)
    stds = params_df["标准差(Std)"].values.astype(float)
    dists = params_df["分布"].tolist()
    dist_params_list = params_df["分布参数"].tolist()

    variances = []
    for i in range(n_params):
        samples_i = generate_sample(dists[i], means[i], stds[i], dist_params_list[i], n_sim)
        results_i = []
        for val in samples_i:
            context_vals = means.copy()
            context_vals[i] = val
            res = safe_eval_with_mapping(formula, param_names, context_vals, param_letters)
            if not np.isnan(res):
                results_i.append(res)
        var_i = np.var(results_i, ddof=1) if len(results_i) > 1 else 0.0
        variances.append(var_i)

    total_var = sum(variances)
    contributions = [v / total_var if total_var > 0 else 0.0 for v in variances]

    df_contrib = pd.DataFrame({
        "参数": param_names,
        "方差贡献": variances,
        "贡献百分比": contributions
    })
    df_contrib = df_contrib.sort_values("贡献百分比", ascending=False).reset_index(drop=True)
    df_contrib["贡献百分比_显示"] = df_contrib["贡献百分比"].apply(lambda x: f"{x:.1%}")
    return df_contrib, contributions, param_names

def plot_pdf(dist: str, mean: float, std: float, dist_params: Dict, ax):
    if dist == "正态分布（完整）":
        x = np.linspace(mean - 4*std, mean + 4*std, 200)
        y = stats.norm.pdf(x, mean, std)
        ax.plot(x, y, 'b-')
        ax.fill_between(x, y, alpha=0.3)
        ax.set_title(f"N(μ={mean:.1f}, σ={std:.2f})", fontsize=8)
    elif dist == "正态分布（正值）":
        a, b = (0 - mean) / std if std > 0 else -np.inf, np.inf
        if std == 0:
            x = [max(mean, 0)]
            y = [1]
        else:
            x = np.linspace(0, mean + 4*std, 200)
            y = stats.truncnorm.pdf(x, a, b, loc=mean, scale=std)
        ax.plot(x, y, 'g-')
        ax.fill_between(x, y, alpha=0.3)
        ax.set_title(f"TruncNorm(≥0)", fontsize=8)
    elif dist == "正态分布（负值）":
        a, b = -np.inf, (0 - mean) / std if std > 0 else np.inf
        if std == 0:
            x = [min(mean, 0)]
            y = [1]
        else:
            x = np.linspace(mean - 4*std, 0, 200)
            y = stats.truncnorm.pdf(x, a, b, loc=mean, scale=std)
        ax.plot(x, y, 'r-')
        ax.fill_between(x, y, alpha=0.3)
        ax.set_title(f"TruncNorm(≤0)", fontsize=8)
    elif dist == "均匀分布":
        low = dist_params.get("low", mean - 3*std)
        high = dist_params.get("high", mean + 3*std)
        x = np.linspace(low, high, 200)
        y = stats.uniform.pdf(x, low, high-low)
        ax.plot(x, y, 'purple')
        ax.fill_between(x, y, alpha=0.3)
        ax.set_title(f"U({low:.1f}, {high:.1f})", fontsize=8)
    elif dist == "对数正态分布":
        mean_log = dist_params.get("mean_log", 0.0)
        sigma_log = dist_params.get("sigma_log", 1.0)
        x = np.linspace(0, np.exp(mean_log + 3*sigma_log), 200)
        y = stats.lognorm.pdf(x, sigma_log, scale=np.exp(mean_log))
        ax.plot(x, y, 'orange')
        ax.fill_between(x, y, alpha=0.3)
        ax.set_title(f"LogN(μlog={mean_log:.1f}, σlog={sigma_log:.2f})", fontsize=8)
    elif dist == "威布尔分布":
        shape = dist_params.get("shape", 1.0)
        scale = dist_params.get("scale", 1.0)
        x = np.linspace(0, scale * 3, 200)
        y = stats.weibull_min.pdf(x, shape, scale=scale)
        ax.plot(x, y, 'brown')
        ax.fill_between(x, y, alpha=0.3)
        ax.set_title(f"Weibull(k={shape:.1f}, λ={scale:.1f})", fontsize=8)
    elif dist == "三角分布":
        left = dist_params.get("left", mean - 3*std)
        mode = dist_params.get("mode", mean)
        right = dist_params.get("right", mean + 3*std)
        x = np.linspace(left, right, 200)
        y = stats.triang.pdf(x, (mode-left)/(right-left), loc=left, scale=right-left)
        ax.plot(x, y, 'olive')
        ax.fill_between(x, y, alpha=0.3)
        ax.set_title(f"Tri({left:.1f}, {mode:.1f}, {right:.1f})", fontsize=8)
    ax.set_xlabel("Value", fontsize=6)
    ax.set_ylabel("Density", fontsize=6)
    ax.tick_params(axis='both', labelsize=6)

def plot_histogram(results, bin_centers, hist_counts, x_pdf, pdf_theory, usl, lsl, output_name, n_sim):
    fig, ax = plt.subplots(figsize=(11, 6), dpi=100)
    ax.bar(bin_centers, hist_counts, width=(bin_centers[1]-bin_centers[0])*0.9, alpha=0.6, label="Histogram", color="#3498db")
    area = np.sum(hist_counts) * (bin_centers[1]-bin_centers[0])
    ax.plot(x_pdf, pdf_theory * area, 'r-', linewidth=2, label="Gaussian Fitting")
    if usl is not None:
        ax.axvline(usl, color='green', linestyle='--', linewidth=1.5, label=f"USL = {usl:.2f}")
    if lsl is not None:
        ax.axvline(lsl, color='orange', linestyle='--', linewidth=1.5, label=f"LSL = {lsl:.2f}")
    stats_text = f"NO.={n_sim}\nAVE={np.mean(results):.2f}\nSTD={np.std(results, ddof=1):.4f}\nMAX={np.max(results):.2f}\nMIN={np.min(results):.2f}"
    ax.text(0.95, 0.95, stats_text, transform=ax.transAxes, fontsize=9, verticalalignment='top', horizontalalignment='right', bbox=dict(boxstyle='round', facecolor='white', alpha=0.8, edgecolor='gray'))
    ax.legend(loc='upper right', bbox_to_anchor=(0.95, 0.72), fontsize=9)
    ax.set_xlabel(output_name, fontsize=11)
    ax.set_ylabel("Frequency", fontsize=11)
    ax.set_title(f"{output_name} Distribution", fontsize=13, fontweight='bold')
    ax.grid(axis='y', linestyle='--', alpha=0.3)
    return fig

def plot_contribution_horizontal(contributions: List[float], param_names: List[str], output_name: str):
    non_zero = [(p, c) for p, c in zip(param_names, contributions) if c > 0]
    if not non_zero:
        fig, ax = plt.subplots(figsize=(8, 4))
        ax.text(0.5, 0.5, "No significant contribution", ha='center', va='center')
        ax.set_title(f"Design Factor Effect % on {output_name}", fontsize=13, fontweight='bold')
        return fig
    names, vals = zip(*non_zero)
    sorted_indices = np.argsort(vals)
    names = [names[i] for i in sorted_indices]
    vals = [vals[i] for i in sorted_indices]
    fig, ax = plt.subplots(figsize=(9, max(4, len(names)*0.4)))
    bars = ax.barh(names, vals, color='#2ecc71')
    for bar, val in zip(bars, vals):
        ax.text(val + 0.01, bar.get_y() + bar.get_height()/2, f'{val:.1%}', va='center', fontsize=9)
    ax.set_xlabel("Effect %", fontsize=11)
    ax.set_title(f"Design Factor Effect % on {output_name}", fontsize=13, fontweight='bold')
    ax.set_xlim(0, max(vals) * 1.15)
    ax.grid(axis='x', linestyle='--', alpha=0.5)
    ax.legend().remove()
    return fig

def compute_cpk_ppm(results: np.ndarray, usl: Optional[float], lsl: Optional[float]):
    mean_out = np.mean(results)
    std_out = np.std(results, ddof=1)
    if std_out == 0:
        cpk = None
    else:
        if usl is not None and lsl is not None:
            cpk = min((usl - mean_out) / (3 * std_out), (mean_out - lsl) / (3 * std_out))
        elif usl is not None:
            cpk = (usl - mean_out) / (3 * std_out)
        elif lsl is not None:
            cpk = (mean_out - lsl) / (3 * std_out)
        else:
            cpk = None
    failures_up = np.sum(results > usl) / len(results) * 1e6 if usl is not None else None
    failures_dn = np.sum(results < lsl) / len(results) * 1e6 if lsl is not None else None
    failures_all = None
    if failures_up is not None and failures_dn is not None:
        failures_all = failures_up + failures_dn
    elif failures_up is not None:
        failures_all = failures_up
    elif failures_dn is not None:
        failures_all = failures_dn
    return cpk, failures_all, failures_up, failures_dn

def generate_word_report(raw, usl, lsl, n_sim, seed, formula, params_df, param_letters, analyst_name, analyst_title, output_name):
    results = raw["results"]
    cpk, failures_all, failures_up, failures_dn = compute_cpk_ppm(results, usl, lsl)

    # 创建 Word 文档
    doc = Document()
    
    # 标题
    title = doc.add_heading(f"{output_name} - DFSS模拟分析报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 分析人信息
    doc.add_heading("分析人信息", level=1)
    p = doc.add_paragraph()
    p.add_run(f"分析人姓名：{analyst_name if analyst_name else '未填写'}").bold = True
    p.add_run(f"\n头衔：{analyst_title if analyst_title else '未填写'}")
    
    # 模拟设置
    doc.add_heading("1. 模拟设置", level=1)
    doc.add_paragraph(f"输出变量名称：{output_name}")
    doc.add_paragraph(f"公式：{formula}")
    doc.add_paragraph(f"模拟次数：{n_sim}")
    doc.add_paragraph(f"随机种子：{seed}")
    doc.add_paragraph(f"规格上限 (USL)：{usl if usl is not None else '无'}")
    doc.add_paragraph(f"规格下限 (LSL)：{lsl if lsl is not None else '无'}")
    
    # 输入参数表
    doc.add_heading("2. 输入参数表", level=1)
    # 将 DataFrame 转换为 Word 表格
    table = doc.add_table(rows=len(params_df)+1, cols=len(params_df.columns))
    table.style = 'Table Grid'
    # 表头
    for j, col in enumerate(params_df.columns):
        table.cell(0, j).text = col
    # 数据
    for i, row in params_df.iterrows():
        for j, col in enumerate(params_df.columns):
            if col == "分布参数":
                # 分布参数是字典，只显示分布类型即可，避免冗长
                table.cell(i+1, j).text = str(row[col]) if row[col] else "{}"
            else:
                table.cell(i+1, j).text = str(row[col])
    
    # 模拟结果统计
    doc.add_heading("3. 模拟结果统计", level=1)
    doc.add_paragraph(f"均值：{raw['mean']:.2f}")
    doc.add_paragraph(f"标准差：{raw['std']:.4f}")
    doc.add_paragraph(f"最大值：{raw['max']:.2f}")
    doc.add_paragraph(f"最小值：{raw['min']:.2f}")
    if cpk is not None:
        doc.add_paragraph(f"Cpk：{cpk:.2f}")
        doc.add_paragraph(f"Failure All (ppm)：{failures_all:.2f}" if failures_all is not None else "Failure All：-")
        doc.add_paragraph(f"Failure Up (ppm)：{failures_up:.2f}" if failures_up is not None else "Failure Up：-")
        doc.add_paragraph(f"Failure Dn (ppm)：{failures_dn:.2f}" if failures_dn is not None else "Failure Dn：-")
    
    # 分布直方图
    doc.add_heading("4. 分布直方图", level=1)
    fig_hist = plot_histogram(results, raw["bin_centers"], raw["hist_counts"], raw["x_pdf"], raw["pdf_theory"], usl, lsl, output_name, n_sim)
    buf_hist = BytesIO()
    fig_hist.savefig(buf_hist, format='png', dpi=150, bbox_inches='tight')
    buf_hist.seek(0)
    doc.add_picture(buf_hist, width=Inches(6))
    plt.close(fig_hist)
    
    # 设计参数影响百分比
    doc.add_heading("5. 设计参数对 " + output_name + " 影响百分比", level=1)
    fig_barh = plot_contribution_horizontal(raw["contributions"], raw["param_names"], output_name)
    buf_barh = BytesIO()
    fig_barh.savefig(buf_barh, format='png', dpi=150, bbox_inches='tight')
    buf_barh.seek(0)
    doc.add_picture(buf_barh, width=Inches(6))
    plt.close(fig_barh)
    
    # 贡献百分比数据表
    doc.add_heading("详细数据表", level=2)
    df_contrib = raw["df_contrib"].copy()
    df_contrib["贡献百分比"] = df_contrib["贡献百分比_显示"]
    contrib_table = doc.add_table(rows=len(df_contrib)+1, cols=2)
    contrib_table.style = 'Table Grid'
    contrib_table.cell(0, 0).text = "参数"
    contrib_table.cell(0, 1).text = "贡献百分比"
    for i, row in df_contrib.iterrows():
        contrib_table.cell(i+1, 0).text = row["参数"]
        contrib_table.cell(i+1, 1).text = row["贡献百分比"]
    
    # 联系信息（页脚）
    doc.add_page_break()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("联系：电邮 Techlife2027@gmail.com").italic = True
    footer.add_run(f"\n报告生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # 保存到 BytesIO
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes

def main():
    st.markdown('<div class="main-title">📊 Para_Variation - 基于蒙特卡洛模拟分析</div>', unsafe_allow_html=True)
    st.markdown("根据输入参数的分布进行随机抽样，计算用户定义的公式结果，分析输出分布及各参数贡献度。")

    with st.sidebar:
        st.markdown("## ⚙️ 模拟设置")
        n_sim = st.number_input("模拟次数 (Trail number)", min_value=100, max_value=100000, value=1000, step=100)
        st.markdown("#### 规格限（可留空）")
        usl_sidebar = st.text_input("规格上限 (USL)", value=st.session_state.usl_str, key="usl_sidebar", on_change=sync_usl_from_sidebar)
        lsl_sidebar = st.text_input("规格下限 (LSL)", value=st.session_state.lsl_str, key="lsl_sidebar", on_change=sync_lsl_from_sidebar)
        st.session_state.usl_str = usl_sidebar
        st.session_state.lsl_str = lsl_sidebar
        seed = st.number_input("随机种子", value=42, step=1)

        # 添加“关于分析系统”说明（润色后）
        st.markdown("---")
        st.markdown("### 关于分析系统")
        st.markdown("""
        - **设计变量**：根据输入参数的公式计算得出。
        - **参数抽样**：每个参数独立根据其概率密度函数（PDF）进行随机抽样。
        
        **输出结果：**
        - 预测设计变量在量产阶段的分布形态、均值及失效率。
        - 量化各设计参数对输出变量的影响百分比。
        - 通过调节规格上下限，快速确定合理的失效率目标，从而定义合理的工程规格。
        """)

        # 分析人信息
        st.markdown("---")
        st.markdown("### 分析人信息")
        analyst_name = st.text_input("分析人姓名", value=st.session_state.analyst_name, key="analyst_name_input")
        analyst_title = st.text_input("分析人头衔（可选）", value=st.session_state.analyst_title, key="analyst_title_input")
        st.session_state.analyst_name = analyst_name
        st.session_state.analyst_title = analyst_title

        # 联系信息
        st.markdown("---")
        st.markdown("**联系：**")
        st.markdown("电邮: Techlife2027@gmail.com")

    # 参数输入表格（与之前相同，略，但完整代码已包含）
    st.markdown('<div class="section-header">📝 参数输入</div>', unsafe_allow_html=True)

    header_cols = st.columns([0.3, 1.5, 1, 1, 1.2, 0.3])
    header_cols[0].markdown("**字母**")
    header_cols[1].markdown("**参数名称**")
    header_cols[2].markdown("**均值(Typ)**")
    header_cols[3].markdown("**标准差(Std)**")
    header_cols[4].markdown("**分布**")
    header_cols[5].markdown("**删除**")

    rows_data = []
    for idx, row in st.session_state.params.iterrows():
        letter = chr(ord('A') + idx)
        cols = st.columns([0.3, 1.5, 1, 1, 1.2, 0.3])
        with cols[0]:
            st.markdown(f'<div class="param-letter">{letter}</div>', unsafe_allow_html=True)
        with cols[1]:
            name = st.text_input("", value=row["参数名称"], key=f"param_name_{idx}", label_visibility="collapsed")
        with cols[2]:
            mean_val = st.number_input("", value=float(row["均值(Typ)"]), step=1.0, key=f"param_mean_{idx}", label_visibility="collapsed")
        with cols[3]:
            std_val = st.number_input("", value=float(row["标准差(Std)"]), step=0.01, format="%.4f", key=f"param_std_{idx}", label_visibility="collapsed")
        with cols[4]:
            dist_val = st.selectbox("", DISTRIBUTIONS, index=DISTRIBUTIONS.index(row["分布"]) if row["分布"] in DISTRIBUTIONS else 0, key=f"param_dist_{idx}", label_visibility="collapsed")
        with cols[5]:
            delete = st.button("🗑️", key=f"del_{idx}")

        current_dist_params = row.get("分布参数", {}) if isinstance(row.get("分布参数"), dict) else {}
        if dist_val in ["均匀分布", "对数正态分布", "威布尔分布", "三角分布"]:
            if dist_val == "均匀分布" and "low" not in current_dist_params:
                current_dist_params["low"] = mean_val - 3 * std_val
                current_dist_params["high"] = mean_val + 3 * std_val
            elif dist_val == "对数正态分布" and "mean_log" not in current_dist_params:
                current_dist_params["mean_log"] = 0.0
                current_dist_params["sigma_log"] = 1.0
            elif dist_val == "威布尔分布" and "shape" not in current_dist_params:
                current_dist_params["shape"] = 1.0
                current_dist_params["scale"] = 1.0
            elif dist_val == "三角分布" and "left" not in current_dist_params:
                current_dist_params["left"] = mean_val - 3 * std_val
                current_dist_params["mode"] = mean_val
                current_dist_params["right"] = mean_val + 3 * std_val

        need_expand = dist_val in ["均匀分布", "对数正态分布", "威布尔分布", "三角分布"]
        if need_expand:
            with st.expander(f"⚙️ 配置 {dist_val} 参数", expanded=True):
                if dist_val == "均匀分布":
                    low = st.number_input("下限", value=float(current_dist_params.get("low", mean_val - 3*std_val)), key=f"uniform_low_{idx}", step=0.1)
                    high = st.number_input("上限", value=float(current_dist_params.get("high", mean_val + 3*std_val)), key=f"uniform_high_{idx}", step=0.1)
                    if low >= high:
                        st.error("下限必须小于上限")
                    else:
                        current_dist_params["low"] = low
                        current_dist_params["high"] = high
                elif dist_val == "对数正态分布":
                    mean_log = st.number_input("对数均值 (μ_log)", value=float(current_dist_params.get("mean_log", 0.0)), key=f"lognorm_meanlog_{idx}", step=0.1)
                    sigma_log = st.number_input("对数标准差 (σ_log)", value=float(current_dist_params.get("sigma_log", 1.0)), key=f"lognorm_sigmalog_{idx}", step=0.05, format="%.3f")
                    if sigma_log <= 0:
                        st.error("对数标准差必须大于0")
                    else:
                        current_dist_params["mean_log"] = mean_log
                        current_dist_params["sigma_log"] = sigma_log
                elif dist_val == "威布尔分布":
                    shape = st.number_input("形状参数 (k)", value=float(current_dist_params.get("shape", 1.0)), key=f"weibull_shape_{idx}", step=0.1, min_value=0.1)
                    scale = st.number_input("尺度参数 (λ)", value=float(current_dist_params.get("scale", 1.0)), key=f"weibull_scale_{idx}", step=0.1, min_value=0.1)
                    if shape <= 0 or scale <= 0:
                        st.error("形状和尺度参数必须 > 0")
                    else:
                        current_dist_params["shape"] = shape
                        current_dist_params["scale"] = scale
                elif dist_val == "三角分布":
                    left = st.number_input("最小值", value=float(current_dist_params.get("left", mean_val - 3*std_val)), key=f"tri_left_{idx}", step=0.1)
                    mode = st.number_input("最可能值", value=float(current_dist_params.get("mode", mean_val)), key=f"tri_mode_{idx}", step=0.1)
                    right = st.number_input("最大值", value=float(current_dist_params.get("right", mean_val + 3*std_val)), key=f"tri_right_{idx}", step=0.1)
                    if not (left <= mode <= right):
                        st.error("必须满足：最小值 ≤ 最可能值 ≤ 最大值")
                    else:
                        current_dist_params["left"] = left
                        current_dist_params["mode"] = mode
                        current_dist_params["right"] = right

                fig, ax = plt.subplots(figsize=(4, 2))
                plot_pdf(dist_val, mean_val, std_val, current_dist_params, ax)
                st.pyplot(fig)
                plt.close(fig)

        rows_data.append((name, mean_val, std_val, dist_val, current_dist_params, delete, letter))

    new_params = []
    for (name, mean_val, std_val, dist_val, dist_params, delete, letter) in rows_data:
        if not delete:
            new_params.append({
                "参数名称": name,
                "均值(Typ)": mean_val,
                "标准差(Std)": std_val,
                "分布": dist_val,
                "分布参数": dist_params
            })
    if st.button("➕ 添加参数行", use_container_width=True):
        new_params.append({
            "参数名称": "新参数",
            "均值(Typ)": 0.0,
            "标准差(Std)": 0.0,
            "分布": "正态分布（完整）",
            "分布参数": {}
        })

    st.session_state.params = pd.DataFrame(new_params)
    update_param_letters()

    # 公式定义区域
    st.markdown('<div class="section-header">📐 公式定义（设计值）</div>', unsafe_allow_html=True)

    st.markdown('<span class="big-label">📌 设计变量名称</span>', unsafe_allow_html=True)
    output_name = st.text_input("", value=st.session_state.output_name, key="output_name_input", label_visibility="collapsed")
    st.session_state.output_name = output_name if output_name.strip() else "Output"

    st.markdown('<span class="big-label">📝 计算公式</span>', unsafe_allow_html=True)
    st.markdown('<div class="formula-hint">💡 可直接在公式中使用字母（A, B, C...）代表对应参数，系统将自动识别。例如：A*E*7/1000*60/(B+C+D)</div>', unsafe_allow_html=True)
    formula = st.text_area("", value=st.session_state.formula, height=100, key="formula_input", label_visibility="collapsed")
    st.session_state.formula = formula
    st.caption("支持的运算: + - * / **, 括号, 函数: sqrt, exp, log, sin, cos, tan, pi, e 等。公式中的空格会被自动优化。")

    design_val = compute_design_value(st.session_state.params, formula, st.session_state.param_letters)
    if design_val is not None and not np.isnan(design_val):
        st.markdown(f"""
        <div class="design-value-card">
            <strong>📌 当前设计值（基于均值）:</strong><br>
            <span class="design-value-number">{output_name} = {design_val:.2f}</span>
        </div>
        """, unsafe_allow_html=True)
    else:
        st.warning("公式无效或参数不匹配，无法计算设计值。请检查公式中的字母是否与上方对应关系一致，并确保运算正确。")

    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        if st.button("开始\n蒙特卡洛模拟", type="primary", use_container_width=True):
            if st.session_state.params.isnull().values.any():
                st.error("参数表中存在空值，请检查！")
                st.stop()
            param_names = st.session_state.params["参数名称"].astype(str).tolist()
            if len(set(param_names)) != len(param_names):
                st.error("参数名称必须唯一！")
                st.stop()
            if not formula.strip():
                st.error("公式不能为空！")
                st.stop()

            with st.spinner("正在进行蒙特卡洛模拟..."):
                sim_res = run_monte_carlo(st.session_state.params, formula, n_sim, st.session_state.param_letters, seed)
            if sim_res is None:
                st.stop()

            with st.spinner("正在计算各参数贡献度..."):
                df_contrib, contributions, param_names = sensitivity_analysis(st.session_state.params, formula, n_sim, st.session_state.param_letters, seed)

            st.session_state.sim_results_raw = {
                "results": sim_res["results"],
                "samples": sim_res["samples"],
                "mean": sim_res["mean"],
                "std": sim_res["std"],
                "max": sim_res["max"],
                "min": sim_res["min"],
                "hist_counts": sim_res["hist_counts"],
                "bin_edges": sim_res["bin_edges"],
                "bin_centers": sim_res["bin_centers"],
                "x_pdf": sim_res["x_pdf"],
                "pdf_theory": sim_res["pdf_theory"],
                "param_names": sim_res["param_names"],
                "df_contrib": df_contrib,
                "contributions": contributions,
                "params_df": st.session_state.params,
                "output_name": output_name,
                "formula": formula,
            }

    if st.session_state.sim_results_raw is not None:
        raw = st.session_state.sim_results_raw
        results = raw["results"]
        output_name = raw["output_name"]
        usl = parse_limit(st.session_state.usl_str)
        lsl = parse_limit(st.session_state.lsl_str)
        cpk, failures_all, failures_up, failures_dn = compute_cpk_ppm(results, usl, lsl)

        st.markdown(f'<div class="section-header">📈 模拟结果: {output_name}</div>', unsafe_allow_html=True)
        col1, col2, col3 = st.columns(3)
        with col1:
            st.markdown(f'<div class="metric-card"><div class="metric-label">{output_name} 均值</div><div class="metric-value">{raw["mean"]:.2f}</div></div>', unsafe_allow_html=True)
            st.markdown(f'<div class="metric-card"><div class="metric-label">{output_name} 标准差</div><div class="metric-value">{raw["std"]:.4f}</div></div>', unsafe_allow_html=True)
        with col2:
            st.markdown(f'<div class="metric-card"><div class="metric-label">最大值</div><div class="metric-value">{raw["max"]:.2f}</div></div>', unsafe_allow_html=True)
            st.markdown(f'<div class="metric-card"><div class="metric-label">最小值</div><div class="metric-value">{raw["min"]:.2f}</div></div>', unsafe_allow_html=True)
        with col3:
            if cpk is not None:
                st.markdown(f'<div class="metric-card"><div class="metric-label">Cpk</div><div class="metric-value">{cpk:.2f}</div></div>', unsafe_allow_html=True)
            else:
                st.markdown('<div class="metric-card"><div class="metric-label">Cpk</div><div class="metric-value">-</div></div>', unsafe_allow_html=True)

        st.markdown('<div class="section-header">Failure ppm level</div>', unsafe_allow_html=True)
        st.caption("💡 可调节上下限以实时观察PPM水平的变化（留空表示无此限）")
        col_left, col_right = st.columns([1, 2])
        with col_left:
            main_usl = st.text_input("规格上限 (USL)", value=st.session_state.usl_str, key="main_usl", on_change=sync_usl_from_main)
            main_lsl = st.text_input("规格下限 (LSL)", value=st.session_state.lsl_str, key="main_lsl", on_change=sync_lsl_from_main)
            st.session_state.usl_str = main_usl
            st.session_state.lsl_str = main_lsl
            usl = parse_limit(main_usl)
            lsl = parse_limit(main_lsl)
            cpk, failures_all, failures_up, failures_dn = compute_cpk_ppm(results, usl, lsl)
        with col_right:
            if cpk is not None:
                def fmt(v): return f"{v:.2f}" if v is not None else "-"
                st.markdown(f"""
                <table class="ppm-table">
                    <tr><th>CPK</th><th>Failure All</th><th>Failure Up</th><th>Failure Dn</th><tr>
                    <tr><td style="text-align:center">{fmt(cpk)}</td><td style="text-align:center">{fmt(failures_all)}</td><td style="text-align:center">{fmt(failures_up)}</td><td style="text-align:center">{fmt(failures_dn)}</td></tr>
                </table>
                """, unsafe_allow_html=True)
            else:
                st.info("未提供任何规格限，无法计算CPK和PPM。")

        # 分布直方图
        st.markdown(f"### {output_name} 分布直方图")
        fig_hist = plot_histogram(results, raw["bin_centers"], raw["hist_counts"], raw["x_pdf"], raw["pdf_theory"], usl, lsl, output_name, n_sim)
        st.pyplot(fig_hist)
        st.caption(f"横轴：{output_name}   |   纵轴：频次")

        # 设计参数影响百分比
        st.markdown(f"### {output_name} 设计参数影响百分比")
        fig_barh = plot_contribution_horizontal(raw["contributions"], raw["param_names"], output_name)
        st.pyplot(fig_barh)
        st.caption("横轴：影响百分比   |   纵轴：设计参数")

        with st.expander("查看贡献百分比数据表"):
            st.dataframe(raw["df_contrib"][["参数", "贡献百分比_显示"]].rename(columns={"贡献百分比_显示": "贡献百分比"}), use_container_width=True)

        with st.expander("查看全部模拟数据"):
            samples_df = pd.DataFrame(raw["samples"], columns=raw["param_names"])
            samples_df[output_name] = results
            st.dataframe(samples_df.round(2), use_container_width=True, height=400)
            csv = samples_df.to_csv(index=False, float_format="%.6f")
            st.download_button("📥 下载模拟数据 (CSV)", data=csv, file_name=f"monte_carlo_data_{output_name}.csv", mime="text/csv")

        # 生成 Word 报告
        doc_bytes = generate_word_report(raw, usl, lsl, n_sim, seed, formula, st.session_state.params, st.session_state.param_letters, st.session_state.analyst_name, st.session_state.analyst_title, output_name)
        st.download_button("📄 下载专业报告 (Word)", data=doc_bytes, file_name=f"DFSS_Report_{output_name}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        st.success("模拟完成！")

if __name__ == "__main__":
    main()
